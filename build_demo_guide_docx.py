"""Build MemoryWatch_Demo_Guide.docx from content."""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)
GREY  = RGBColor(80, 80, 80)
CODE_BG = "F0F0F0"

doc = Document()
sec = doc.sections[0]
sec.page_width  = Cm(21)
sec.page_height = Cm(29.7)
for a in ('left_margin','right_margin','top_margin','bottom_margin'):
    setattr(sec, a, Cm(2.5))

doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.color.rgb = BLACK

def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(20); r.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(6)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(14); r.font.color.rgb = BLACK
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(12); r.font.color.rgb = BLACK
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)

TOKEN_RE = re.compile(r'(\*\*[^*]+?\*\*|`[^`]+?`)')

def parse_inline(line):
    segs = []
    pos = 0
    for m in TOKEN_RE.finditer(line):
        if m.start() > pos:
            segs.append((line[pos:m.start()], False, False))
        tok = m.group(0)
        if tok.startswith('**'):
            segs.append((tok[2:-2], True, False))
        else:
            segs.append((tok[1:-1], False, True))
        pos = m.end()
    if pos < len(line):
        segs.append((line[pos:], False, False))
    return segs

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(indent)
    for txt, bold, code in parse_inline(text):
        r = p.add_run(txt)
        r.font.name = 'Courier New' if code else 'Arial'
        r.font.size  = Pt(10.5 if code else 11)
        r.font.color.rgb = BLACK
        r.bold = bold
    return p

def numbered(text, n):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.8)
    r0 = p.add_run(f"{n}. ")
    r0.bold = True; r0.font.name = 'Arial'; r0.font.size = Pt(11); r0.font.color.rgb = BLACK
    for txt, bold, code in parse_inline(text):
        r = p.add_run(txt)
        r.font.name = 'Courier New' if code else 'Arial'
        r.font.size  = Pt(10.5 if code else 11)
        r.font.color.rgb = BLACK
        r.bold = bold

def bullet(text, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(indent + 0.5)
    r0 = p.add_run(u'•  ')
    r0.font.name = 'Arial'; r0.font.size = Pt(11); r0.font.color.rgb = BLACK
    for txt, bold, code in parse_inline(text):
        r = p.add_run(txt)
        r.font.name = 'Courier New' if code else 'Arial'
        r.font.size  = Pt(10.5 if code else 11)
        r.font.color.rgb = BLACK
        r.bold = bold

def code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Courier New'; r.font.size = Pt(9.5); r.font.color.rgb = BLACK
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), CODE_BG)
        p.paragraph_format.element.get_or_add_pPr().append(shading)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_table(headers, rows, col_widths):
    from docx.shared import Cm as C
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn as Q
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    # header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.width = C(col_widths[i])
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(10); r.font.color.rgb = BLACK
    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri+1].cells[ci]
            cell.width = C(col_widths[ci])
            p = cell.paragraphs[0]
            for txt, bold, code in parse_inline(val):
                r = p.add_run(txt)
                r.font.name = 'Courier New' if code else 'Arial'
                r.font.size  = Pt(9.5 if code else 10)
                r.font.color.rgb = BLACK
                r.bold = bold
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─────────────────────────────────────────────
# DOCUMENT CONTENT
# ─────────────────────────────────────────────

title("MemoryWatch — Demo Guide")
body("A step-by-step guide for running the MemoryWatch live memory monitoring demo. No prior setup needed beyond Docker.")
divider()

# ── What is MemoryWatch? ──
h1("What is MemoryWatch?")
body("MemoryWatch is a memory-based intrusion detection system. It watches running processes on a Linux machine and flags any that behave suspiciously — for example, a process that suddenly opens anonymous executable memory regions (a sign of shellcode injection) or one that reads another process's memory (credential scraping).")
body("It works in two stages:")
bullet("**Baseline** — watches the live system for a short period and learns what normal looks like")
bullet("**Live monitoring** — keeps watching and scores every process against that baseline. Anything that deviates enough triggers an alert.")
body("The detection algorithm is an **Isolation Forest** — a machine learning model that learns normal behaviour and flags outliers, without ever needing labelled examples of attacks.")

divider()

# ── Files included ──
h1("Files Included")
add_table(
    ["File", "What it does"],
    [
        ["`utils.py`",                   "The Isolation Forest algorithm, built from scratch in NumPy"],
        ["`memory_features.py`",          "Reads live process data from /proc (memory usage, file descriptors, executable memory regions, syscall patterns, I/O activity)"],
        ["`live_memory_monitor.py`",      "Runs the full pipeline: baseline → train model → live scoring → alerts"],
        ["`demo_workload.py`",            "A harmless test process that mimics a memory attack signature (anonymous executable memory + repetitive syscalls) so you can see the monitor catch something"],
        ["`generate_memory_dataset.py`",  "Generates a realistic synthetic dataset of normal and attack process profiles"],
        ["`test_on_memory_dataset.py`",   "Evaluates the detection pipeline on that synthetic dataset and prints detection rates per attack type"],
    ],
    [4.5, 11.5]
)

divider()

# ── Requirements ──
h1("Requirements")
bullet("**Docker Desktop** — download free from docker.com/products/docker-desktop")
bullet("That's it. Everything else runs inside the container.")

divider()

# ── Setup ──
h1("Setup (one time only)")

h2("1. Install Docker Desktop and open it")
body("You will see a whale icon in your menu bar when it is running.")

h2("2. Put all 6 project files in one folder")
body("For example, create a folder on your Desktop:")
code_block(["~/Desktop/memorywatch/", "  utils.py", "  memory_features.py", "  live_memory_monitor.py",
            "  demo_workload.py", "  generate_memory_dataset.py", "  test_on_memory_dataset.py"])

h2("3. Open Terminal and start a Linux container")
body("Run this command — it mounts your folder directly inside the container so files are immediately available:")
code_block(["docker run -it -v ~/Desktop/memorywatch:/workspace ubuntu:22.04 bash"])
body("This downloads Ubuntu (~70MB, first time only) and drops you into a Linux terminal. Your files are at /workspace.")

h2("4. Install dependencies (inside the container)")
code_block(["apt update && apt install -y python3 python3-numpy python3-pandas"])

divider()

# ── Test 1 ──
h1("Test 1 — Synthetic Dataset Evaluation")
body("This test validates that the detection pipeline works correctly on a realistic simulated dataset before touching a live system.")
code_block(["cd /workspace", "python3 generate_memory_dataset.py", "python3 test_on_memory_dataset.py"])

h2("What generate_memory_dataset.py does")
body("Creates memory_dataset.csv with 698 rows of simulated process data — 600 normal rows (systemd, Chrome, postgres, sshd, bash, python3, vim) plus 98 attack rows across three types:")
bullet("`heap_spray` — many anonymous executable memory regions, very low syscall entropy (repetitive loop)")
bullet("`proc_scraper` — holds an open file descriptor into another process's memory, massive read I/O")
bullet("`mem_injection` — sudden large memory spike, rapid growth of executable regions")

h2("What test_on_memory_dataset.py does")
body("Trains the Isolation Forest on the normal rows only, then scores all rows and reports overall metrics (accuracy, precision, recall, FPR, F1, AUROC), per-attack-type detection rates, and which features deviated most for each attack type.")

h2("Expected output")
code_block([
    "AUROC: 0.9594   <- excellent ranking ability",
    "heap_spray:     0.0%   <- stealthy, overlaps with normal JIT processes",
    "mem_injection: 14.3%   <- vmrss_delta spike catches some",
    "proc_scraper:   3.6%   <- conservative threshold misses most",
])
body("The low detection rates mirror what happened in the main project — the model is tuned for very low false alarms (FPR ~1%), which means it only alerts when highly confident. The AUROC of 0.9594 confirms it correctly ranks anomalies above normal traffic.")

divider()

# ── Test 2 ──
h1("Test 2 — Live Real-Time Monitoring")
body("This runs the actual monitor against live processes in the container.")
code_block(["python3 live_memory_monitor.py --threshold-pct 95"])

h2("What happens step by step")
numbered("**Baseline phase (12s)** — polls all running processes every second, builds a feature matrix, trains the Isolation Forest on it, sets the anomaly threshold at the 95th percentile of normal scores", 1)
numbered("**Live phase (25s)** — keeps polling and scoring every second", 2)
numbered("At the 8-second mark, automatically spawns `demo_workload.py`", 3)
numbered("`demo_workload.py` creates 4 anonymous executable memory pages (simulating shellcode staging) and runs a tight repetitive write loop (simulating low-entropy heap spray behaviour)", 4)
numbered("If the monitor flags it, you will see an `[ALERT]` line with the PID, process name, score, and which features triggered it", 5)
numbered("At the end: `demo_workload.py was DETECTED` or `MISSED`", 6)

h2("Note on Docker vs. a real VM")
body("In Docker the container only has 2–3 processes, so the baseline is thin and the model sometimes misses the demo workload (both are Python processes and look similar). On a real Linux machine with 100+ diverse processes running, the baseline is much richer and detection is more reliable. The `--threshold-pct 95` flag makes the model more sensitive to compensate.")

divider()

# ── Features ──
h1("What the Features Mean")
add_table(
    ["Feature", "What it measures", "Why it matters"],
    [
        ["`vmrss_kb`",               "How much RAM the process is using",              "Baseline for memory behaviour"],
        ["`vmrss_delta`",            "Change in RAM since last poll",                  "Sudden spike = memory injection"],
        ["`fd_count`",               "Number of open file descriptors",                "Unusual count = suspicious activity"],
        ["`fd_proc_mem_open`",       "Whether this process has a handle into another process's memory", "Direct credential-scraping indicator"],
        ["`maps_exec_region_count`", "Number of executable memory regions",            "Normal processes have a fixed set from their binary and libraries"],
        ["`maps_anonymous_exec`",    "Anonymous and executable memory regions",        "Should be near 0; shellcode lives here"],
        ["`exec_region_growth_rate`","How fast new executable regions appear",         "Heap spray creates many in a short time"],
        ["`syscall_entropy`",        "How varied the process's syscall pattern is",    "Low entropy = repetitive loop = suspicious"],
        ["`io_read_bytes_delta`",    "Bytes read since last poll",                     "A proc scraper reads gigabytes scanning victim memory"],
        ["`io_write_bytes_delta`",   "Bytes written since last poll",                  "Unusually high = exfiltration-like behaviour"],
    ],
    [4.0, 5.5, 6.5]
)

divider()

# ── Tuning ──
h1("Tuning the Sensitivity")
code_block([
    "# Default — very conservative, almost no false alarms",
    "python3 live_memory_monitor.py",
    "",
    "# More sensitive — catches more, may flag some normal processes",
    "python3 live_memory_monitor.py --threshold-pct 95",
    "",
    "# Even more sensitive",
    "python3 live_memory_monitor.py --threshold-pct 90",
    "",
    "# Longer baseline for richer normal profile",
    "python3 live_memory_monitor.py --baseline-duration 30 --threshold-pct 95",
])

divider()

# ── Limitations ──
h1("Limitations")
bullet("**Requires Linux** — /proc is Linux-only. Docker provides this; macOS or Windows alone cannot run the live monitor.")
bullet("**Sparse container baseline** — Docker containers have very few processes, so the model has limited data to define normal. A real Linux VM gives much better results.")
bullet("**No root access** — features like counting ptrace attach calls require kernel-level tracing (eBPF/auditd) that needs root. The current monitor approximates these using unprivileged /proc reads.")
bullet("**Syscall entropy is approximated** — the monitor samples /proc/[pid]/syscall once per polling interval rather than recording every syscall, because full tracing would require root.")

divider()

# ── Quick reference ──
h1("Quick Reference")
code_block([
    "# Start fresh container (run from Mac Terminal)",
    "docker run -it -v ~/Desktop/memorywatch:/workspace ubuntu:22.04 bash",
    "",
    "# Inside container — first time setup",
    "apt update && apt install -y python3 python3-numpy python3-pandas",
    "",
    "# Run synthetic dataset test",
    "cd /workspace",
    "python3 generate_memory_dataset.py",
    "python3 test_on_memory_dataset.py",
    "",
    "# Run live monitor",
    "python3 live_memory_monitor.py --threshold-pct 95",
])

# ── Save ──
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'MemoryWatch_Demo_Guide.docx')
doc.save(out)

import docx as _docx
d2 = _docx.Document(out)
zoom_els = d2.settings.element.findall(qn('w:zoom'))
for z in zoom_els:
    z.getparent().remove(z)
d2.save(out)
print(f"Saved: {out}")
